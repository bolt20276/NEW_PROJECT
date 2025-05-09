
import streamlit as st
import pandas as pd
import re
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize,SyllableTokenizer
from nltk.stem import WordNetLemmatizer
import datetime
import email
import spacy
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import imaplib
import email.utils
import os
import base64
import tempfile
import PyPDF2
import docx
import io
from email.header import decode_header
import time
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
import io

def download_nltk_resources():
    """Download required NLTK resources with proper error handling."""
    st.set_page_config(page_title="CV Parser & Ranker", layout="wide")
    resources = ['punkt', 'stopwords', 'wordnet']
    for resource in resources:
        try:
            nltk.data.find(f'tokenizers/{resource}')
            st.write(f"✓ NLTK resource '{resource}' already downloaded")
        except LookupError:
            st.write(f"Downloading NLTK resource '{resource}'...")
            nltk.download(resource, quiet=True)
            st.write(f"✓ NLTK resource '{resource}' downloaded successfully")

# Ensure NLTK resources are available
try:
    download_nltk_resources()
except Exception as e:
    st.error(f"Error setting up NLTK resources: {str(e)}")
    st.info("Please try running this command in your terminal: python -m nltk.downloader punkt stopwords wordnet")

# Initialize stop words and lemmatizer after ensuring resources are available
try:
    stop_words = set(stopwords.words('english'))
    lemmatizer = WordNetLemmatizer()
except Exception as e:
    st.error(f"Error initializing NLTK components: {str(e)}")
    stop_words = set()  # Fallback to empty set
    lemmatizer = None   # Will need to handle this in functions that use lemmatizer

# Load spaCy model with proper error handling
try:
    nlp = spacy.load("en_core_web_sm")
except:
    st.warning("First-time setup: Downloading spaCy model...")
    import os
    os.system("python -m spacy download en_core_web_sm")
    try:
        nlp = spacy.load("en_core_web_sm")
    except Exception as e:
        st.error(f"Error loading spaCy model: {str(e)}")
        st.info("Please try running this command in your terminal: python -m spacy download en_core_web_sm")
        # Define a minimal fallback function if spaCy fails
        class MinimalNLP:
            def __call__(self, text):
                class MinimalDoc:
                    ents = []
                return MinimalDoc()
        nlp = MinimalNLP()


def debug_info(message, data=None):
    """Helper function to print debug information"""
    print(f"DEBUG: {message}")
    if data is not None:
        print(f"DATA: {data}")

def test_email_connection(email_provider, email_address, password):
    """Test email server connection without fetching CVs"""
    try:
        if email_provider == "Gmail":
            server = "imap.gmail.com"
        else:  # Yahoo
            server = "imap.mail.yahoo.com"
        
        mail = imaplib.IMAP4_SSL(server)
        mail.login(email_address, password)
        mail.select("inbox")
        
        # Get mailbox info
        status, data = mail.status('INBOX', '(MESSAGES)')
        if status == 'OK':
            message_count = int(re.search(r'MESSAGES\s+(\d+)', data[0].decode()).group(1))
            
           
            status, email_ids = mail.search(None, 'ALL')
            if status == 'OK':
                email_id_list = email_ids[0].split()
                if email_id_list:
                    # Get the 5 most recent emails (or fewer if there aren't that many)
                    recent_emails = email_id_list[-min(5, len(email_id_list)):]
                    subjects = []
                    
                    for email_id in recent_emails:
                        status, msg_data = mail.fetch(email_id, '(BODY[HEADER.FIELDS (SUBJECT)])')
                        if status == 'OK':
                            subject_data = msg_data[0][1].decode()
                            subject = re.search(r'Subject: (.*)', subject_data)
                            if subject:
                                subjects.append(subject.group(1))
                    
                    mail.logout()
                    return True, f"Connected successfully. {message_count} messages in inbox. Recent subjects: {', '.join(subjects)}"
            
            mail.logout()
            return True, f"Connected successfully. {message_count} messages in inbox."
        
        mail.logout()
        return True, "Connected successfully but couldn't get message count."
    
    except Exception as e:
        return False, f"Connection failed: {str(e)}"

def add_logging_to_streamlit():
    """Add logging to display debug information in the Streamlit UI"""
    import logging
    
    # Create a custom handler that will store logs in the session state
    class SessionStateHandler(logging.Handler):
        def emit(self, record):
            log_entry = self.format(record)
            if 'log_messages' not in st.session_state:
                st.session_state.log_messages = []
            st.session_state.log_messages.append(log_entry)
    
    # Configure the root logger
    logger = logging.getLogger()
    logger.setLevel(logging.INFO)
    
    # Add our custom handler
    handler = SessionStateHandler()
    formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
    handler.setFormatter(formatter)
    logger.addHandler(handler)
    
    return logger

# Function to display debug logs in the UI
def display_debug_logs():
    if 'log_messages' in st.session_state and st.session_state.log_messages:
        with st.expander("Debug Logs"):
            for log in st.session_state.log_messages:
                st.text(log)
    
    # Option to clear logs
    if st.button("Clear Logs"):
        st.session_state.log_messages = []
        st.rerun()


def preprocess_text(text):
    """Basic text preprocessing with error handling"""
    if not text:
        return ""
    
    try:
        # Convert to lowercase
        text = text.lower()
        tk = SyllableTokenizer()
        
        # Remove special characters and numbers
        text = re.sub(r'[^\w\s]', ' ', text)
        text = re.sub(r'\d+', ' ', text)
        
        # Tokenize with error handling
        try:
            tokens = tk.tokenize(text)
        except LookupError:
            # Fallback simple tokenization if NLTK fails
            st.warning("NLTK tokenization failed. Using simple tokenization instead.")
            nltk.download('punkt')  # Try downloading again
            try:
                tokens = tk.tokenize(text)
            except:
                # Very basic fallback tokenization
                tokens = text.split()
        
        # Remove stopwords and lemmatize with error handling
        processed_tokens = []
        for token in tokens:
            if token not in stop_words:
                try:
                    if lemmatizer:
                        processed_token = lemmatizer.lemmatize(token)
                    else:
                        processed_token = token
                    processed_tokens.append(processed_token)
                except:
                    processed_tokens.append(token)
        
        return ' '.join(processed_tokens)
    
    except Exception as e:
        st.error(f"Error in text preprocessing: {str(e)}")
        # Return original text if processing fails
        return text.lower()
    
def extract_email_address(text):
    """Extract email addresses from text"""
    email_pattern = r'[\w\.-]+@[\w\.-]+\.\w+'
    emails = re.findall(email_pattern, text)
    return emails[0] if emails else ""

def extract_name(text):
    """Extract person name using spaCy's NER"""
    doc = nlp(text[:1000])  # Process first 1000 chars for efficiency
    for ent in doc.ents:
        if ent.label_ == "PERSON":
            return ent.text
    # Fallback: Look for common resume header patterns
    name_patterns = [
        r'Name\s*:\s*([\w\s]+)',
        r'^([\w\s]+)$',  # First line of CV often contains name
        r'Resume of ([\w\s]+)'
    ]
    for pattern in name_patterns:
        matches = re.search(pattern, text[:500], re.MULTILINE | re.IGNORECASE)
        if matches:
            return matches.group(1).strip()
    return "Unknown"

def extract_years_of_experience(text):
    """Extract years of experience from CV"""
    experience_patterns = [
        r'(\d+)\+?\s*(?:years|yrs)(?:\s*of)?\s*experience',
        r'experience\s*(?:of)?\s*(\d+)\+?\s*(?:years|yrs)',
        r'(?:professional|work|industry)\s*experience\s*(?:of)?\s*(\d+)\+?\s*(?:years|yrs)',
        r'worked for (\d+)\+?\s*(?:years|yrs)'
    ]
    
    for pattern in experience_patterns:
        matches = re.search(pattern, text, re.IGNORECASE)
        if matches:
            try:
                return int(matches.group(1))
            except:
                pass
    
    # If no direct mention of years, try to calculate from employment history
    try:
        # Look for date ranges in employment history
        date_ranges = re.findall(r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\d{4})\s*(?:-|to|–)\s*((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\d{4}|Present|Current|Now)', text, re.IGNORECASE)
        if date_ranges:
            total_months = 0
            current_year = datetime.datetime.now().year
            current_month = datetime.datetime.now().month
            
            for start, end in date_ranges:
                # Parse start date
                if re.search(r'\d{4}', start):
                    start_year = int(re.search(r'\d{4}', start).group())
                    start_month_match = re.search(r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*', start, re.IGNORECASE)
                    start_month = datetime.datetime.strptime(start_month_match.group(1)[:3], '%b').month if start_month_match else 1
                else:
                    continue
                
                # Parse end date
                if re.search(r'\d{4}', end):
                    end_year = int(re.search(r'\d{4}', end).group())
                    end_month_match = re.search(r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*', end, re.IGNORECASE)
                    end_month = datetime.datetime.strptime(end_month_match.group(1)[:3], '%b').month if end_month_match else 12
                elif re.search(r'Present|Current|Now', end, re.IGNORECASE):
                    end_year = current_year
                    end_month = current_month
                else:
                    continue
                
                months_worked = (end_year - start_year) * 12 + (end_month - start_month)
                if months_worked > 0:
                    total_months += months_worked
            
            return round(total_months / 12, 1)  # Convert months to years
    except:
        pass
    
    return 0  # Default if no experience information found


def extract_skills(cv_text, skill_list):
    """Extract skills from CV based on a predefined skill list with error handling"""
    if not cv_text or not skill_list:
        return []
    
    try:
        found_skills = []
        processed_cv = preprocess_text(cv_text)
        
        for skill in skill_list:
            try:
                # Create a regex pattern to match the skill as a whole word
                pattern = r'\b' + re.escape(skill.lower()) + r'\b'
                if re.search(pattern, processed_cv):
                    found_skills.append(skill)
            except:
                # If there's an error with a specific skill, just skip it
                continue
        
        return found_skills
    except Exception as e:
        st.error(f"Error extracting skills: {str(e)}")
        return []  # Return empty list if extraction fails

def extract_certifications(text):
    """Extract certifications from CV"""
    certification_patterns = [
        r'(?:certification|certificate|certified|certify|diploma)(?:\s*in)?\s*([\w\s]+)',
        r'([\w\s]+)\s*certification',
        r'((?:AWS|Azure|Google Cloud|CISSP|PMP|Agile|Scrum|ITIL|CompTIA|Cisco|Microsoft|CISA|CISM|CFA|CPA|PMI|CAPM|Six Sigma|TOGAF|CCNA|MCSA|MCSE|MCITP|RHCE|RHCSA|Oracle|IBM|SAP|Salesforce)(?:\s*[\w\s]+)?)'
    ]
    
    certifications = []
    for pattern in certification_patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for match in matches:
            cert = match.group(1).strip()
            # Filter out common false positives
            if len(cert) > 3 and cert.lower() not in ['the', 'a', 'an', 'that', 'this', 'these', 'those']:
                certifications.append(cert)
    
    # Deduplicate preserving order
    seen = set()
    unique_certs = []
    for cert in certifications:
        if cert.lower() not in seen:
            seen.add(cert.lower())
            unique_certs.append(cert)
    
    return unique_certs

def extract_education(text):
    """Extract education information from CV with precise details"""
    education_pattern = r'''
        (?:  # Degree types
            B\.?A\.?|B\.?S\.?|B\.?Sc\.?|Bachelor|
            M\.?A\.?|M\.?S\.?|M\.?Sc\.?|Master|MBA|
            Ph\.?D\.?|Doctorate|DPhil|
            Associate|Diploma|Certificate
        )
        # FIXED GROUP BELOW (added missing parenthesis)
        (?:\s+(?:of\s+(?:Science|Arts)|in|degree))?  # Degree descriptors
        \s*(?:in|,)?\s*
        ([\w\s&]+?)  # Field of study
        (?:\s*(?:from|at|,)\s*([-\w\s&]+))?  # Institution
        (?:\s*\(?(\d{4})\s*[-–]\s*(\d{4}|Present)\)?)?  # Dates
    '''

    education = []
    seen = set()
    
    for match in re.finditer(education_pattern, text, re.VERBOSE | re.IGNORECASE):
        degree_type = match.group(0).split()[0].strip()
        field = match.group(1).strip() if match.group(1) else ''
        institution = match.group(2).strip() if match.group(2) else ''
        start_year = match.group(3) if match.group(3) else ''
        end_year = match.group(4) if match.group(4) else ''

        # Filter out false positives and format the entry
        if institution and any(x in institution.lower() for x in ['university', 'college', 'institute', 'school']):
            entry = {
                'degree': degree_type,
                'field': field,
                'institution': institution,
                'years': f"{start_year}-{end_year}" if start_year else ''
            }
            entry_str = f"{entry['degree']} in {entry['field']}, {entry['institution']} {entry['years']}".strip()
            
            if entry_str.lower() not in seen:
                seen.add(entry_str.lower())
                education.append(entry_str)

    # Fallback pattern for entries without clear institution names
    if not education:
        fallback_pattern = r'''
            (?:Education|Academic\s+Background)[\s:-]+
            ((?:.|\n)+?)(?=\n\s*\n|Work\s+Experience|Employment|Skills|$)
        '''
        education_section = re.search(fallback_pattern, text, re.VERBOSE | re.IGNORECASE)
        if education_section:
            education = [line.strip() for line in education_section.group(1).split('\n') if line.strip()]

    return education

def extract_text_from_pdf(file_content):
    """Extract text from PDF file content with improved error handling"""
    try:
        with io.BytesIO(file_content) as pdf_file:
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            if pdf_reader.is_encrypted:
                return ""
            
            text = ""
            for page_num in range(len(pdf_reader.pages)):
                try:
                    # Extract text from each page individually to isolate potential errors
                    page_text = pdf_reader.pages[page_num].extract_text()
                    if page_text:
                        text += page_text + " "
                except Exception as e:
                    print(f"Error extracting text from PDF page {page_num}: {str(e)}")
                    # Continue to next page if one fails
                    continue
            return text
    except Exception as e:
        print(f"Error extracting text from PDF: {str(e)}")
        return ""

def extract_text_from_docx(file_content):
    """Extract text from DOCX file content with better error handling"""
    try:
        with io.BytesIO(file_content) as docx_file:
            doc = docx.Document(docx_file)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            # Also extract text from tables if present
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            return text
    except Exception as e:
        print(f"Error extracting text from DOCX: {str(e)}")
        return ""

def extract_text_from_file(file_content, file_name):
    """Extract text from a file based on its extension with improved error handling"""
    if not file_content:
        return ""
        
    file_extension = os.path.splitext(file_name)[1].lower()
    try:
        if file_extension == '.pdf':
            return extract_text_from_pdf(file_content)
        elif file_extension in ['.docx', '.doc']:
            return extract_text_from_docx(file_content)
        elif file_extension == '.txt':
            return file_content.decode('utf-8', errors='ignore')
        else:
            print(f"Unsupported file extension: {file_extension}")
            return ""
    except Exception as e:
        print(f"Error in extract_text_from_file: {str(e)}")
        # Try a generic approach as fallback
        try:
            return file_content.decode('utf-8', errors='ignore')
        except:
            return ""


def decode_email_subject(subject):
    """Decode email subject that might be encoded"""
    decoded_subject = ""
    if subject:
        decoded_chunks = decode_header(subject)
        for chunk, encoding in decoded_chunks:
            if isinstance(chunk, bytes):
                if encoding:
                    decoded_subject += chunk.decode(encoding, errors='replace')
                else:
                    decoded_subject += chunk.decode('utf-8', errors='replace')
            else:
                decoded_subject += chunk
    return decoded_subject



def fetch_emails_gmail(email_address, password, search_criteria, time_frame_days=30, max_emails=50):
    """Fetch emails from Gmail using IMAP"""
    cv_texts = []
    cv_filenames = []
    cv_sources = []
    
    try:
        # Connect to Gmail's IMAP server
        mail = imaplib.IMAP4_SSL("imap.gmail.com")
        mail.login(email_address, password)
        mail.select("inbox")
        
        # Calculate the date threshold
        date_threshold = (datetime.datetime.now() - datetime.timedelta(days=time_frame_days)).strftime("%d-%b-%Y")
        
        # Create search criteria
        search_string = f'(SINCE {date_threshold})'
        if search_criteria:
            search_string = f'(SINCE {date_threshold} SUBJECT "{search_criteria}")'
        
        # Search for emails
        status, email_ids = mail.search(None, search_string)
        
        if status != 'OK':
            return False, [], [], []
        
        # Get the list of email IDs
        email_id_list = email_ids[0].split()
        
        # Limit the number of emails to process
        if len(email_id_list) > max_emails:
            email_id_list = email_id_list[-max_emails:]
        
        # Process each email
        for email_id in reversed(email_id_list):  # Process newest first
            status, msg_data = mail.fetch(email_id, '(RFC822)')
            
            if status != 'OK':
                continue
            
            raw_email = msg_data[0][1]
            msg = email.message_from_bytes(raw_email)
            
            # Get email details
            subject = decode_email_subject(msg.get("Subject", ""))
            from_address = msg.get("From", "")
            date_tuple = email.utils.parsedate_tz(msg.get("Date", ""))
            
            if date_tuple:
                email_date = datetime.datetime.fromtimestamp(email.utils.mktime_tz(date_tuple))
                days_old = (datetime.datetime.now() - email_date).days
                
                # Skip if email is too old
                if days_old > time_frame_days:
                    continue
                
                email_date_str = email_date.strftime("%Y-%m-%d %H:%M:%S")
            else:
                email_date_str = "Unknown date"
            
            # Flag to track if we found attachments in this email
            found_attachments = False
            
            # Process attachments
            if msg.is_multipart():
                for part in msg.walk():
                    content_disposition = str(part.get("Content-Disposition", ""))
                    
                    # Check if it's an attachment
                    if "attachment" in content_disposition or "filename" in content_disposition:
                        filename = part.get_filename()
                        
                        if filename:
                            # Decode filename if needed
                            if isinstance(filename, str):
                                filename_parts = decode_header(filename)
                                filename = ""
                                for chunk, encoding in filename_parts:
                                    if isinstance(chunk, bytes):
                                        if encoding:
                                            filename += chunk.decode(encoding, errors='replace')
                                        else:
                                            filename += chunk.decode('utf-8', errors='replace')
                                    else:
                                        filename += chunk
                            
                            # Check if it's a CV/resume file
                            file_ext = os.path.splitext(filename.lower())[1]
                            if file_ext in ['.pdf', '.doc', '.docx', '.txt']:
                                payload = part.get_payload(decode=True)
                                if payload:
                                    # Extract text from file
                                    extracted_text = extract_text_from_file(payload, filename)
                                    
                                    if extracted_text and len(extracted_text) > 200:  # Only if substantial text
                                        cv_texts.append(extracted_text)
                                        cv_filenames.append(filename)
                                        cv_sources.append(f"Email from: {from_address} | Subject: {subject} | Date: {email_date_str}")
                                        found_attachments = True
            
            # Try to extract CV from email body itself if no attachments found in this email
            # or if the subject suggests it contains a CV/resume
            if (not found_attachments or 
                "resume" in subject.lower() or 
                "cv" in subject.lower() or
                "application" in subject.lower()):
                
                body = ""
                
                if msg.is_multipart():
                    for part in msg.walk():
                        content_type = part.get_content_type()
                        if content_type == "text/plain":
                            payload = part.get_payload(decode=True)
                            if payload:
                                body += payload.decode('utf-8', errors='ignore')
                else:
                    payload = msg.get_payload(decode=True)
                    if payload:
                        body = payload.decode('utf-8', errors='ignore')
                
                if body and len(body) > 500:  # Only consider substantial text
                    cv_texts.append(body)
                    cv_filenames.append(f"Email body - {subject}")
                    cv_sources.append(f"Email from: {from_address} | Subject: {subject} | Date: {email_date_str}")
        
        # Logout
        mail.logout()
        
        # Debug information
        print(f"Found {len(cv_texts)} potential CVs")
        
        return True, cv_texts, cv_filenames, cv_sources
    
    except Exception as e:
        print(f"Error in fetch_emails_gmail: {str(e)}")
        return False, [], [], []

def fetch_emails_yahoo(email_address, password, search_criteria, time_frame_days=30, max_emails=50):
    """Fetch emails from Yahoo Mail using IMAP"""
    cv_texts = []
    cv_filenames = []
    cv_sources = []
    
    try:
        # Connect to Yahoo's IMAP server
        mail = imaplib.IMAP4_SSL("imap.mail.yahoo.com")
        mail.login(email_address, password)
        mail.select("inbox")
        
        # Calculate the date threshold
        date_threshold = (datetime.datetime.now() - datetime.timedelta(days=time_frame_days)).strftime("%d-%b-%Y")
        
        # Create search criteria
        search_string = f'(SINCE {date_threshold})'
        if search_criteria:
            search_string = f'(SINCE {date_threshold} SUBJECT "{search_criteria}")'
        
        # Search for emails
        status, email_ids = mail.search(None, search_string)
        
        if status != 'OK':
            return False, [], [], []
        
        # Get the list of email IDs
        email_id_list = email_ids[0].split()
        
        # Limit the number of emails to process
        if len(email_id_list) > max_emails:
            email_id_list = email_id_list[-max_emails:]
        
        # Process each email
        for email_id in reversed(email_id_list):  # Process newest first
            status, msg_data = mail.fetch(email_id, '(RFC822)')
            
            if status != 'OK':
                continue
            
            raw_email = msg_data[0][1]
            msg = email.message_from_bytes(raw_email)
            
            # Get email details
            subject = decode_email_subject(msg.get("Subject", ""))
            from_address = msg.get("From", "")
            date_tuple = email.utils.parsedate_tz(msg.get("Date", ""))
            
            if date_tuple:
                email_date = datetime.datetime.fromtimestamp(email.utils.mktime_tz(date_tuple))
                days_old = (datetime.datetime.now() - email_date).days
                
                # Skip if email is too old
                if days_old > time_frame_days:
                    continue
                
                email_date_str = email_date.strftime("%Y-%m-%d %H:%M:%S")
            else:
                email_date_str = "Unknown date"
            
          
            found_attachments = False
            
            if msg.is_multipart():
                for part in msg.walk():
                    content_disposition = str(part.get("Content-Disposition", ""))
                    
                    # Check if it's an attachment
                    if "attachment" in content_disposition or "filename" in content_disposition:
                        filename = part.get_filename()
                        
                        if filename:
                            # Decode filename if needed
                            if isinstance(filename, str):
                                filename_parts = decode_header(filename)
                                filename = ""
                                for chunk, encoding in filename_parts:
                                    if isinstance(chunk, bytes):
                                        if encoding:
                                            filename += chunk.decode(encoding, errors='replace')
                                        else:
                                            filename += chunk.decode('utf-8', errors='replace')
                                    else:
                                        filename += chunk
                            
                            # Check if it's a CV/resume file
                            file_ext = os.path.splitext(filename.lower())[1]
                            if file_ext in ['.pdf', '.doc', '.docx', '.txt']:
                                payload = part.get_payload(decode=True)
                                if payload:
                                    # Extract text from file
                                    extracted_text = extract_text_from_file(payload, filename)
                                    
                                    if extracted_text and len(extracted_text) > 200:  # Only if substantial text
                                        cv_texts.append(extracted_text)
                                        cv_filenames.append(filename)
                                        cv_sources.append(f"Email from: {from_address} | Subject: {subject} | Date: {email_date_str}")
                                        found_attachments = True
            
            # Try to extract CV from email body itself if no attachments found in this email
            # or if the subject suggests it contains a CV/resume
            if (not found_attachments or 
                "resume" in subject.lower() or 
                "cv" in subject.lower() or
                "application" in subject.lower()):
                
                body = ""
                
                if msg.is_multipart():
                    for part in msg.walk():
                        content_type = part.get_content_type()
                        if content_type == "text/plain":
                            payload = part.get_payload(decode=True)
                            if payload:
                                body += payload.decode('utf-8', errors='ignore')
                else:
                    payload = msg.get_payload(decode=True)
                    if payload:
                        body = payload.decode('utf-8', errors='ignore')
                
                if body and len(body) > 500:  # Only consider substantial text
                    cv_texts.append(body)
                    cv_filenames.append(f"Email body - {subject}")
                    cv_sources.append(f"Email from: {from_address} | Subject: {subject} | Date: {email_date_str}")
        
        # Logout
        mail.logout()
        
        # Debug information
        print(f"Found {len(cv_texts)} potential CVs")
        
        return True, cv_texts, cv_filenames, cv_sources
    
    except Exception as e:
        print(f"Error in fetch_emails_yahoo: {str(e)}")
        return False, [], [], []
    
def parse_email_content(email_text, time_frame_days=30):
    """Parse email content directly from text"""
    if not email_text:
        return []
    
    # Parse the email text
    parsed_email = email.message_from_string(email_text)
    
    # Extract email date
    date_str = parsed_email.get("Date", "")
    email_date = None
    if date_str:
        try:
            # Parse the email date (simplified for demo)
            email_date = datetime.datetime.strptime(date_str, "%a, %d %b %Y %H:%M:%S %z")
        except:
            try:
                # Try an alternative format
                email_date = datetime.datetime.strptime(date_str, "%d %b %Y %H:%M:%S %z")
            except:
                email_date = datetime.datetime.now()  # Fallback to current date
    else:
        email_date = datetime.datetime.now()
    
    # Check if the email is within the timeframe
    current_date = datetime.datetime.now().replace(tzinfo=email_date.tzinfo)
    days_old = (current_date - email_date).days
    
    if days_old > time_frame_days:
        return []
    
    # Extract CV texts
    cv_texts = []
    
    # Get email body
    if parsed_email.is_multipart():
        for part in parsed_email.get_payload():
            content_type = part.get_content_type()
            if content_type == "text/plain":
                cv_texts.append(part.get_payload())
    else:
        cv_texts.append(parsed_email.get_payload())
    
    all_cvs = []
    for text in cv_texts:
        potential_cvs = re.split(r'\n{4,}', text)
        for potential_cv in potential_cvs:
            if len(potential_cv.strip()) > 200:  # Only consider substantial text blocks
                all_cvs.append(potential_cv)
    
    return all_cvs

def parse_jd_from_file(uploaded_file):
    """Parse Job Description from an uploaded file (PDF, DOCX, TXT)"""
    if uploaded_file is None:
        return ""
    
    try:
        file_extension = os.path.splitext(uploaded_file.name)[1].lower()
        
        # Get file content
        file_content = uploaded_file.read()
        
        # Extract text based on file type
        if file_extension == '.pdf':
            return extract_text_from_pdf(file_content)
        elif file_extension in ['.docx', '.doc']:
            return extract_text_from_docx(file_content)
        elif file_extension == '.txt':
            return file_content.decode('utf-8', errors='ignore')
        else:
            st.error(f"Unsupported file format: {file_extension}. Please upload a PDF, DOCX, or TXT file.")
            return ""
    except Exception as e:
        st.error(f"Error parsing file: {str(e)}")
        return ""

def extract_skills_from_jd(jd_text):
    """Extract required skills from job description"""
    # Common technical skills to look for
    common_skills = [
        "Python", "Java", "JavaScript", "TypeScript", "C++", "C#", "Go", "Rust", "Swift",
        "PHP", "Ruby", "Kotlin", "R", "MATLAB", "SQL", "NoSQL", "MongoDB", "MySQL",
        "PostgreSQL", "Oracle", "HTML", "CSS", "React", "Angular", "Vue.js", "Node.js",
        "Django", "Flask", "Spring", "ASP.NET", "Express.js", "TensorFlow", "PyTorch",
        "Keras", "Scikit-learn", "Pandas", "NumPy", "AWS", "Azure", "Google Cloud",
        "Docker", "Kubernetes", "Jenkins", "Git", "SVN", "CI/CD", "DevOps", "Agile",
        "Scrum", "Kanban", "REST API", "GraphQL", "Microservices", "Serverless",
        "Linux", "Windows", "MacOS", "iOS", "Android", "Unity", "Hadoop", "Spark",
        "Kafka", "Redis", "Elasticsearch", "Tableau", "Power BI", "JIRA", "Confluence",
        "Firebase", "Redux", "jQuery", "Bootstrap", "SASS", "LESS", "WebSocket",
        "Blockchain", "Ethereum", "Solidity", "WebRTC", "SEO", "Wordpress", "Shopify",
        "Magento", "Big Data", "Data Science", "Machine Learning", "Artificial Intelligence",
        "NLP", "Computer Vision", "AR/VR", "UI/UX", "Photoshop", "Illustrator", "Figma",
        "Sketch", "InVision", "Penetration Testing", "Cybersecurity", "Ethical Hacking",
        "Network Security", "SIEM", "Splunk", "ELK Stack", "Wireshark", "Selenium",
        "JUnit", "TestNG", "Mocha", "Jasmine", "Jest", "Cypress", "Postman", "SoapUI"
    ]
    
    # Find skills mentioned in the JD
    skills_in_jd = []
    for skill in common_skills:
        # Create pattern to match whole word
        pattern = r'\b' + re.escape(skill) + r'\b'
        if re.search(pattern, jd_text, re.IGNORECASE):
            skills_in_jd.append(skill)
    
    # Also look for skill sections
    skill_section_pattern = r'(?:skills|requirements|qualifications)(?:\s*required)?(?:\s*:)?\s*([\w\s,\.\+\-\&\/]+)'
    matches = re.search(skill_section_pattern, jd_text, re.IGNORECASE)
    if matches:
        skill_section = matches.group(1)
        # Extract additional skills from the skill section
        additional_skills = re.findall(r'\b([A-Za-z][\w\+\#\.]*(?:\s*[\w\+\#\.]+)?)\b', skill_section)
        for skill in additional_skills:
            skill = skill.strip()
            if len(skill) > 2 and skill.lower() not in stop_words and skill not in skills_in_jd:
                skills_in_jd.append(skill)
    
    return skills_in_jd

def score_candidate(cv_data, jd_data):
    """Score candidate based on JD requirements"""
    score = 0
    scoring_components = {}
    
    # 1. Years of experience (up to 30 points)
    yoe = cv_data.get("years_of_experience", 0)
    yoe_score = min(30, yoe * 5)  # 5 points per year, max 30
    score += yoe_score
    scoring_components["Experience"] = {"score": yoe_score, "details": f"{yoe} years of experience"}
    
    # 2. Skills match (up to 40 points)
    required_skills = jd_data.get("skills", [])
    candidate_skills = cv_data.get("skills", [])
    
    if required_skills:
        skills_match_percentage = len(candidate_skills) / len(required_skills) * 100
        skills_score = min(40, skills_match_percentage * 0.4)  # 0.4 points per % match, max 40
        score += skills_score
        scoring_components["Skills"] = {
            "score": skills_score, 
            "details": f"{len(candidate_skills)}/{len(required_skills)} skills matched ({skills_match_percentage:.1f}%)"
        }
    
    # 3. Certifications (up to 15 points)
    certifications = cv_data.get("certifications", [])
    cert_score = min(15, len(certifications) * 5)  # 5 points per certification, max 15
    score += cert_score
    scoring_components["Certifications"] = {
        "score": cert_score, 
        "details": f"{len(certifications)} relevant certifications found"
    }
    
    # 4. Education (up to 15 points)
    education = cv_data.get("education", [])
    edu_score = min(15, len(education) * 5)  # 5 points per education entry, max 15
    score += edu_score
    scoring_components["Education"] = {
        "score": edu_score, 
        "details": f"{len(education)} education credentials found"
    }
    
    # 5. Content similarity using TF-IDF (up to 30 points)
    cv_text = cv_data.get("text", "")
    jd_text = jd_data.get("text", "")
    
    if cv_text and jd_text:
        tfidf_vectorizer = TfidfVectorizer(stop_words='english')
        try:
            tfidf_matrix = tfidf_vectorizer.fit_transform([jd_text, cv_text])
            similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
            similarity_score = min(30, similarity * 30)  # Scale similarity to 30 points max
            score += similarity_score
            scoring_components["Content Similarity"] = {
                "score": similarity_score, 
                "details": f"{similarity:.2f} cosine similarity between CV and JD"
            }
        except:
            # In case of TF-IDF failure, allocate average points
            similarity_score = 15
            score += similarity_score
            scoring_components["Content Similarity"] = {
                "score": similarity_score, 
                "details": "Average score (TF-IDF analysis failed)"
            }
    
    # Total score
    total_score = min(100, score)  # Cap at 100 points
    
    return {
        "total_score": total_score,
        "components": scoring_components
    }

# Initialize session state
def init_session_state():
    """Initialize session state variables"""
    if 'logged_in' not in st.session_state:
        st.session_state.logged_in = False
    if 'email_provider' not in st.session_state:
        st.session_state.email_provider = None
    if 'email_address' not in st.session_state:
        st.session_state.email_address = ""
    if 'cv_data' not in st.session_state:
        st.session_state.cv_data = []
    if 'jd_data' not in st.session_state:
        st.session_state.jd_data = {}
    if 'ranked_candidates' not in st.session_state:
        st.session_state.ranked_candidates = []
    if 'jd_submitted' not in st.session_state:
        st.session_state.jd_submitted = False
    if 'parsing_complete' not in st.session_state:
        st.session_state.parsing_complete = False


def main():
    # st.set_page_config(page_title="CV Parser & Ranker", layout="wide")
    
    # Initialize session state
    init_session_state()
    
    # App title
    st.title("CV Parser & Ranker")
    
    # Navigation tabs - only show all tabs if logged in
    if st.session_state.logged_in:
        tabs = st.tabs(["Email Login", "Job Description", "Parse & Rank Candidates", "Results"])
        login_tab, jd_tab, parse_tab, results_tab = tabs
    else:
        tabs = st.tabs(["Email Login"])
        login_tab = tabs[0]
    
    # Email Login Tab
    with login_tab:
        st.header("Email Login")
    
        if not st.session_state.logged_in:
            st.info("Please login to your email to start fetching CVs")
            
            col1, col2 = st.columns(2)
            
            with col1:
                email_provider = st.selectbox("Email Provider", ["Gmail", "Yahoo Mail"])
                # Store inputs directly in session state
                email_address = st.text_input("Email Address", key="email_input")
                password = st.text_input("Password", type="password", key="password_input")
                
            with col2:
                search_criteria = st.text_input("Search Criteria (Optional)", 
                                            placeholder="e.g., CV, Resume, Application")
                time_frame = st.slider("Time Frame (Days)", min_value=1, max_value=90, value=30)
                max_emails = st.slider("Maximum Emails to Process", min_value=5, max_value=100, value=50)
            
            # Debug information - can be removed in production
            if st.checkbox("Show debug info"):
                st.write(f"Email input length: {len(st.session_state.email_input)}")
                st.write(f"Password input length: {len(st.session_state.password_input)}")
            
            if st.button("Login & Fetch CVs"):
                # Access values directly from session state
                if not st.session_state.email_input or not st.session_state.password_input:
                    st.error("Please enter both email address and password")
                else:
                    with st.spinner("Logging in and fetching CVs..."):
                        try:
                            if email_provider == "Gmail":
                                success, cv_texts, cv_filenames, cv_sources = fetch_emails_gmail(
                                    st.session_state.email_input, st.session_state.password_input, 
                                    search_criteria, time_frame, max_emails)
                            else:  # Yahoo Mail
                                success, cv_texts, cv_filenames, cv_sources = fetch_emails_yahoo(
                                    st.session_state.email_input, st.session_state.password_input, 
                                    search_criteria, time_frame, max_emails)
                            
                            if success:
                                # Display debug info
                                st.write(f"Debug: Found {len(cv_texts)} CVs, {len(cv_filenames)} filenames, {len(cv_sources)} sources")
                                
                                # Process CVs
                                processed_cvs = []
                                
                                with st.spinner(f"Processing {len(cv_texts)} CVs..."):
                                    for i, (cv_text, filename, source) in enumerate(zip(cv_texts, cv_filenames, cv_sources)):
                                        # Extract data from CV
                                        name = extract_name(cv_text)
                                        email = extract_email_address(cv_text)
                                        years_exp = extract_years_of_experience(cv_text)
                                        
                                        # Add to processed CVs
                                        processed_cvs.append({
                                            "id": i,
                                            "name": name,
                                            "email": email,
                                            "years_of_experience": years_exp,
                                            "text": cv_text,
                                            "filename": filename,
                                            "source": source,
                                            "certifications": extract_certifications(cv_text),
                                            "education": extract_education(cv_text)
                                        })
                                
                                st.session_state.cv_data = processed_cvs
                                st.session_state.logged_in = True
                                st.session_state.email_provider = email_provider
                                st.session_state.email_address = st.session_state.email_input
                                
                                st.success(f"Successfully fetched {len(processed_cvs)} CVs! Please go to the Job Description tab.")
                                st.rerun()
                            else:
                                st.error("Failed to login or fetch emails. Please check your credentials and try again.")
                        except Exception as e:
                            st.error(f"Error during email fetching: {str(e)}")
        else:
            st.success(f"Logged in as {st.session_state.email_address} ({st.session_state.email_provider})")
            st.info(f"Found {len(st.session_state.cv_data)} potential CV documents")
            
            # Display brief summary of fetched CVs
            if len(st.session_state.cv_data) > 0:
                with st.expander("View fetched CV summary"):
                    for i, cv in enumerate(st.session_state.cv_data):
                        st.write(f"{i+1}. {cv['name']} - {cv['filename']}")
            
            if st.button("Logout"):
                # Clear session state
                st.session_state.logged_in = False
                st.session_state.email_provider = None
                st.session_state.email_address = ""
                st.session_state.cv_data = []
                st.session_state.jd_data = {}
                st.session_state.ranked_candidates = []
                st.session_state.jd_submitted = False
                st.session_state.parsing_complete = False
                st.rerun()

    if st.session_state.logged_in:
    # Tab 2: CV Source

        with jd_tab:
            st.header("Job Description")
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.subheader("Enter Job Description")
                jd_text = st.text_area("Paste Job Description", height=300)
                
            with col2:
                st.subheader("Or Upload JD File")
                uploaded_file = st.file_uploader("Upload Job Description (PDF, DOCX, TXT)", 
                                                type=["pdf", "docx", "txt"])
                
                if uploaded_file is not None:
                    with st.spinner("Extracting text from file..."):
                        jd_text = parse_jd_from_file(uploaded_file)
            
            if st.button("Submit Job Description"):
                if not jd_text or len(jd_text) < 100:
                    st.error("Please enter a valid job description (minimum 100 characters)")
                else:
                    # Extract skills from JD
                    skills = extract_skills_from_jd(jd_text)
                    
                    # Store JD data
                    st.session_state.jd_data = {
                        "text": jd_text,
                        "skills": skills
                    }
                    
                    st.session_state.jd_submitted = True
                    st.success(f"Job Description processed successfully! Identified {len(skills)} required skills.")
                    st.info("Please proceed to the 'Parse & Rank Candidates' tab.")

        # Parse & Rank Tab  
        with parse_tab:
            st.header("Parse & Rank Candidates")
            
            if not st.session_state.jd_submitted:
                st.warning("Please submit a Job Description first before ranking candidates.")
            else:
                st.info(f"Ready to rank {len(st.session_state.cv_data)} candidates against the submitted Job Description")
                
                if st.button("Parse & Rank Candidates"):
                    if len(st.session_state.cv_data) == 0:
                        st.error("No CVs found to rank. Please check your email settings.")
                    else:
                        with st.spinner("Ranking candidates..."):
                            ranked_candidates = []
                            
                            # Get skills from JD
                            jd_skills = st.session_state.jd_data.get("skills", [])
                            
                            # Process each CV
                            for cv in st.session_state.cv_data:
                                # Extract skills now that we have JD skills
                                cv_skills = extract_skills(cv["text"], jd_skills)
                                cv["skills"] = cv_skills
                                
                                # Score the candidate
                                score_results = score_candidate(cv, st.session_state.jd_data)
                                
                                ranked_candidates.append({
                                    **cv,
                                    "score": score_results["total_score"],
                                    "score_components": score_results["components"]
                                })
                            
                            # Sort by score (descending)
                            ranked_candidates.sort(key=lambda x: x["score"], reverse=True)
                            
                            # Store results
                            st.session_state.ranked_candidates = ranked_candidates
                            st.session_state.parsing_complete = True
                            
                            st.success("Candidates ranked successfully! Please view results in the Results tab.")

upload_tab, analysis_tab, results_tab = st.tabs(["Upload", "Analysis", "Results"])

# Results Tab
with results_tab:
    st.header("Ranked Candidates")
    
    if not st.session_state.parsing_complete:
        st.warning("Please complete the parsing and ranking process first.")
    else:
        # Display top candidates
        st.subheader(f"Top Candidates (Total: {len(st.session_state.ranked_candidates)})")
        
        # Create dataframe of candidates for overview
        candidates_overview = []
        for candidate in st.session_state.ranked_candidates:
            candidates_overview.append({
                "Name": candidate["name"],
                "Email": candidate["email"],
                "Experience (Years)": candidate["years_of_experience"],
                "Skills Match": f"{len(candidate['skills'])}/{len(st.session_state.jd_data.get('skills', []))}",
                "Score": f"{candidate['score']:.1f}/100"
            })
        
        # Display as dataframe
        df = pd.DataFrame(candidates_overview)
        st.dataframe(df, use_container_width=True)

        # PDF Download Section
        if st.button("📄 Generate PDF Report"):
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.lib import colors
            from reportlab.lib.units import inch
            import io
            from datetime import datetime

            # Create PDF buffer
            pdf_buffer = io.BytesIO()

            # Create PDF document
            doc = SimpleDocTemplate(pdf_buffer, pagesize=letter)
            elements = []
            styles = getSampleStyleSheet()

            # Add title
            title_text = "Candidate Ranking Report"
            elements.append(Paragraph(title_text, styles['Title']))
            elements.append(Spacer(1, 0.2*inch))

            # Add report date
            report_date = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            elements.append(Paragraph(f"Generated on: {report_date}", styles['Normal']))
            elements.append(Spacer(1, 0.3*inch))

            # Create data for table
            data = [["Rank", "Name", "Email", "Experience", "Skills Match", "Score"]]
            for idx, candidate in enumerate(st.session_state.ranked_candidates, 1):
                data.append([
                    str(idx),
                    candidate['name'],
                    candidate['email'],
                    str(candidate['years_of_experience']),
                    f"{len(candidate['skills'])}/{len(st.session_state.jd_data.get('skills', []))}",
                    f"{candidate['score']:.1f}/100"
                ])

            # Create table
            table = Table(data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#2A4C7D")),
                ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
                ('ALIGN', (0,0), (-1,-1), 'CENTER'),
                ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                ('FONTSIZE', (0,0), (-1,0), 10),
                ('BOTTOMPADDING', (0,0), (-1,0), 12),
                ('BACKGROUND', (0,1), (-1,-1), colors.HexColor("#F0F2F6")),
                ('GRID', (0,0), (-1,-1), 0.5, colors.black),
                ('FONTSIZE', (0,1), (-1,-1), 9)
            ]))
            
            elements.append(table)
            elements.append(Spacer(1, 0.5*inch))

            # Add summary statistics
            elements.append(Paragraph("Summary Statistics:", styles['Heading2']))
            summary_data = [
                ["Total Candidates", len(st.session_state.ranked_candidates)],
                ["Average Score", f"{df['Score'].str.replace('/100', '').astype(float).mean():.1f}/100"],
                ["Highest Score", f"{df['Score'].str.replace('/100', '').astype(float).max():.1f}/100"],
                ["Lowest Score", f"{df['Score'].str.replace('/100', '').astype(float).min():.1f}/100"]
            ]
            summary_table = Table(summary_data, colWidths=[2.5*inch, 2.5*inch])
            summary_table.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#4B86B4")),
                ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
                ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                ('BACKGROUND', (0,1), (-1,-1), colors.HexColor("#E1E8F0")),
                ('GRID', (0,0), (-1,-1), 0.5, colors.black)
            ]))
            elements.append(summary_table)

            # Build PDF
            doc.build(elements)

            # Create download link
            st.download_button(
                label="⬇️ Download Full Report",
                data=pdf_buffer.getvalue(),
                file_name=f"candidate_ranking_report_{datetime.now().strftime('%Y%m%d')}.pdf",
                mime="application/pdf",
                help="Download detailed candidate ranking report in PDF format"
            )

        # Detailed view of candidates
        st.subheader("Detailed Candidate Information")
        
        # Store selected candidate in session state
        if 'selected_candidate' not in st.session_state:
            st.session_state.selected_candidate = None
        
        # List all candidates as buttons/expanders
        for i, candidate in enumerate(st.session_state.ranked_candidates):
            if st.button(f"{i+1}. {candidate['name']} - Score: {candidate['score']:.1f}/100"):
                st.session_state.selected_candidate = i
        
        # Display selected candidate details outside of any expander
        if st.session_state.selected_candidate is not None:
            candidate = st.session_state.ranked_candidates[st.session_state.selected_candidate]
            
            st.markdown("### Candidate Details")
            col1, col2 = st.columns([1, 2])
            
            with col1:
                st.markdown(f"**Source:** {candidate['source']}")
                st.markdown(f"**File:** {candidate['filename']}")
                st.markdown(f"**Email:** {candidate['email']}")
                st.markdown(f"**Experience:** {candidate['years_of_experience']} years")
                
                # Score breakdown
                st.subheader("Score Breakdown")
                for component, details in candidate['score_components'].items():
                    st.markdown(f"**{component}:** {details['score']:.1f} - {details['details']}")
            
            with col2:
                # Skills match
                st.subheader("Skills Match")
                all_skills = st.session_state.jd_data.get("skills", [])
                matched_skills = candidate["skills"]
                
                for skill in all_skills:
                    if skill in matched_skills:
                        st.markdown(f"✅ {skill}")
                    else:
                        st.markdown(f"❌ {skill}")
                
                # Certifications
                if candidate["certifications"]:
                    st.subheader("Certifications")
                    for cert in candidate["certifications"]:
                        st.markdown(f"• {cert}")
                
                # Education
                if candidate["education"]:
                    st.subheader("Education")
                    for edu in candidate["education"]:
                        st.markdown(f"• {edu}")
            
            # View full CV text
            show_cv = st.checkbox("Show Full CV Text")
            if show_cv:
                st.text_area("CV Text", candidate["text"], height=300)

# Run the application
if __name__ == "__main__":
    main()
