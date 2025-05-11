# Refactored: Resume Analyzer with improved accuracy, modularity, and performance
# Author: Intern (with help from ChatGPT)
# Date: 2025-05

import streamlit as st
import re
import os
import io
import base64
import email
import imaplib
import datetime
import hashlib
import docx
import PyPDF2
import pandas as pd
import spacy
from fpdf import FPDF
from email.header import decode_header
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from rapidfuzz import fuzz

# Load NLP model
try:
    nlp = spacy.load("en_core_web_sm")
except:
    import subprocess
    subprocess.run(["python", "-m", "spacy", "download", "en_core_web_sm"])
    nlp = spacy.load("en_core_web_sm")

# --- Utilities ---
def decode_subject(subject):
    decoded = decode_header(subject)
    return ''.join([part.decode(enc or 'utf-8') if isinstance(part, bytes) else part for part, enc in decoded])

def decode_filename(filename):
    decoded = decode_header(filename)
    return ''.join([part.decode(enc or 'utf-8') if isinstance(part, bytes) else part for part, enc in decoded])

def extract_text_from_pdf(file_content):
    try:
        reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        return " ".join([page.extract_text() or "" for page in reader.pages])
    except:
        return ""

def extract_text_from_docx(file_content):
    try:
        doc = docx.Document(io.BytesIO(file_content))
        return " ".join(p.text for p in doc.paragraphs)
    except:
        return ""

def clean_text(text):
    text = text.replace('\n', ' ').replace('\r', ' ')
    return re.sub(r'[^\x00-\x7F]+', ' ', text)

def hash_text(text):
    return hashlib.sha256(text.encode('utf-8')).hexdigest()

# --- CV Extraction ---
def extract_responsibilities(jd_text):
    """Extract responsibilities or requirements from job description text."""
    matches = re.findall(r'(responsibilities|requirements|key tasks)[:\-\s]*([\s\S]+?)(?=(\n\n|\n[A-Z]))', jd_text, re.I)
    if matches:
        return [line.strip() for m in matches for line in m[1].split("\n") if 5 < len(line.strip()) < 150]
    return []

def extract_name(text):
    doc = nlp(text[:1000])
    for ent in doc.ents:
        if ent.label_ == "PERSON":
            return ent.text
    return "Unknown"

def extract_email(text):
    match = re.search(r'[\w\.-]+@[\w\.-]+', text)
    return match.group(0) if match else ""

def extract_skills(text, jd_text=""):
    """Dynamic skill extraction using both CV and JD"""
    jd_skills = re.findall(r'\b([A-Z][a-z]{3,})\b', jd_text)  # Extract proper nouns from JD
    base_skills = ["Python", "SQL", "AWS", "Azure", "Java", "Spark", 
                  "Hadoop", "Hive", "Kafka", "Airflow", "Kubernetes"]
    all_keywords = list(set(base_skills + jd_skills))
    
    text = text.lower()
    found_skills = []
    for skill in all_keywords:
        if skill.lower() in text:
            found_skills.append(skill)
    return sorted(found_skills, key=lambda x: len(x), reverse=True)[:10]  # Top 10 relevant

def extract_certifications(text):
    """Improved certification extraction"""
    patterns = [
        r'(AWS Certified [\w-]+)|(Certified [\w-]+)|(Lean Six Sigma [\w ]+)',
        r'\b(PM[AP]|CISSP|CCNA|Azure Fundamentals|Google Cloud)\b'
    ]
    certs = set()
    for pattern in patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for match in matches:
            cert = match.group().strip()
            if 3 < len(cert) < 50:
                certs.add(cert)
    return sorted(certs)

def extract_education(text):
    """Enhanced education parsing with NLP"""
    doc = nlp(text)
    education = []
    
    # Pattern matching for degrees
    degree_pattern = r'\b(B\.?S\.?|BSc|Bachelor|M\.?S\.?|MSc|Master|PhD)\b'
    for sent in doc.sents:
        if re.search(degree_pattern, sent.text, re.I):
            education.append(re.sub(r'\s+', ' ', sent.text.strip()))
    
    return sorted(list(set(education))[:3])  # Return max 3 most relevant

def extract_experience(text):
    """Improved experience calculation with months handling"""
    current_year = datetime.datetime.now().year
    current_month = datetime.datetime.now().month
    experience_months = 0
    
    # Find all employment periods
    matches = re.findall(
        r'(\d{4})\s*-\s*(present|current|\d{4})', 
        text, 
        re.IGNORECASE
    )
    
    for start, end in matches:
        start_year = int(start)
        end_year = current_year if end.lower() in ['present', 'current'] else int(end)
        
        # Calculate months
        if end.lower() in ['present', 'current']:
            experience_months += (current_year - start_year) * 12 + current_month
        else:
            experience_months += (end_year - start_year) * 12
            
    return round(experience_months / 12, 1)

def match_responsibilities(cv_text, jd_requirements):
    """Fuzzy match responsibilities using rapidfuzz"""
    matched = []
    cv_sentences = [sent.strip() for sent in re.split(r'[.•]', cv_text) if sent.strip()]
    
    for req in jd_requirements:
        for sent in cv_sentences:
            if fuzz.partial_ratio(req.lower(), sent.lower()) > 85:
                matched.append(sent)
                break
    return list(set(matched))[:5]  # Return max 5 best matches
    
# --- PDF generation ---
def export_cv_to_pdf(candidate):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 8, txt=f"Resume: {candidate['Name']}\n\n{candidate['RawText']}")
    path = f"{candidate['Name'].replace(' ', '_')}_resume.pdf"
    pdf.output(path)
    return path

# --- Email Functions ---
def fetch_attachments(email_user, email_pass, keyword_filter, days_back, max_emails):
    try:
        server = "imap.gmail.com" if "gmail" in email_user else "imap.mail.yahoo.com" if "yahoo" in email_user else "imap-mail.outlook.com"
        mail = imaplib.IMAP4_SSL(server)
        mail.login(email_user, email_pass)
        mail.select("inbox")
        date_filter = (datetime.datetime.now() - datetime.timedelta(days=days_back)).strftime("%d-%b-%Y")
        search = f'(SINCE {date_filter}' + (f' SUBJECT "{keyword_filter}"' if keyword_filter else '') + ')'
        status, messages = mail.search(None, search)
        ids = list(set(messages[0].split()))[-max_emails:]
        attachments = []
        for i in ids[::-1]:
            _, msg_data = mail.fetch(i, "(RFC822)")
            msg = email.message_from_bytes(msg_data[0][1])
            subject = decode_subject(msg.get("Subject", ""))
            for part in msg.walk():
                filename = part.get_filename()
                if filename:
                    filename = decode_filename(filename)
                    ext = os.path.splitext(filename)[1].lower()
                    content = part.get_payload(decode=True)
                    if ext in [".pdf", ".docx"]:
                        attachments.append((filename, content, subject))
        return attachments
    except Exception as e:
        st.error(f"Email error: {e}")
        return []

# --- Scoring ---
def score_cv(cv_text, jd_text):
    try:
        vectorizer = TfidfVectorizer(stop_words='english')
        tfidf = vectorizer.fit_transform([jd_text, cv_text])
        similarity = cosine_similarity(tfidf[0:1], tfidf[1:2])[0][0]
    except:
        similarity = 0
    return round(similarity * 100)

# --- Streamlit UI ---
def main_page():
    st.title("📄 AI Resume Analyzer")
    email_user = st.text_input("Email Address")
    email_pass = st.text_input("Email Password (App Password)", type="password")
    subject_filter = st.text_input("Filter by Subject or Keyword (Optional)")
    days_back = st.slider("Only check emails from the past X days", min_value=1, max_value=90, value=30)
    max_emails = st.slider("Max emails to fetch", min_value=10, max_value=100, value=50)

    st.markdown("### Job Description")
    jd_text = st.text_area("Paste Job Description", height=200)
    uploaded_jd = st.file_uploader("Or upload Job Description (PDF/DOCX/TXT)", type=["pdf", "docx", "txt"])
    
    if uploaded_jd is not None:
        file_content = uploaded_jd.read()
        if uploaded_jd.name.lower().endswith(".pdf"):
            jd_text = extract_text_from_pdf(file_content)
        elif uploaded_jd.name.lower().endswith(".docx"):
            jd_text = extract_text_from_docx(file_content)
        elif uploaded_jd.name.lower().endswith(".txt"):
            jd_text = file_content.decode("utf-8", errors='ignore')
        else:
            st.error("Unsupported file type")
            jd_text = ""

    if st.button("Fetch & Analyze Resumes"):
        if not email_user or not email_pass or not jd_text.strip():
            st.error("Fill all fields.")
        else:
            with st.spinner("Processing resumes..."):
                seen_hashes = set()
                results = []
                attachments = fetch_attachments(email_user, email_pass, subject_filter, days_back, max_emails)
                
                for fname, content, subject in attachments:
                    ext = os.path.splitext(fname)[1].lower()
                    cv_text = extract_text_from_pdf(content) if ext == ".pdf" else extract_text_from_docx(content)
                    cv_text = clean_text(cv_text)
                    if len(cv_text) < 100:
                        continue
                    hash_val = hash_text(cv_text)
                    if hash_val in seen_hashes:
                        continue
                    seen_hashes.add(hash_val)

                    score = score_cv(cv_text, jd_text)
                    name = extract_name(cv_text)
                    email_addr = extract_email(cv_text)
                    skills = extract_skills(cv_text)
                    certs = extract_certifications(cv_text)
                    edu = extract_education(cv_text)
                    exp = extract_experience(cv_text)
                    jd_responsibilities = extract_responsibilities(jd_text)
                    matched = [r for r in jd_responsibilities if r.lower() in cv_text.lower()]

                    pdf_path = export_cv_to_pdf({"Name": name, "RawText": cv_text})
                    with open(pdf_path, "rb") as f:
                        encoded = base64.b64encode(f.read()).decode()
                    download_link = f'<a href="data:application/pdf;base64,{encoded}" download="{pdf_path}">Download PDF</a>'

                    results.append({
                        "Name": name,
                        "Email": email_addr,
                        "Score": score,
                        "Experience": exp,
                        "Skills": ", ".join(skills),
                        "Certifications": ", ".join(certs),
                        "Education": ", ".join(edu),
                        "Responsibilities Matched": ", ".join(matched),
                        "File": download_link,
                        "Subject": subject,
                    })

                if results:
                    df = pd.DataFrame(results)
                    # Store results in session state
                    st.session_state.results_df = df
                    st.session_state.show_results = True
                    st.rerun()
                else:
                    st.warning("No valid resumes found.")

def results_page():
    st.title("📊 Analysis Results")
    
    if 'results_df' not in st.session_state:
        st.warning("No results found. Please analyze resumes first.")
        if st.button("Back to Main Page"):
            st.session_state.show_results = False
            st.rerun()
        return

    df = st.session_state.results_df
    
    # Add interactive filters
    st.sidebar.header("Filters")
    min_score = st.sidebar.slider("Minimum Score", 0, 100, 0)
    min_experience = st.sidebar.slider("Minimum Experience (years)", 0, 50, 0)
    
    filtered_df = df[
        (df['Score'] >= min_score) &
        (df['Experience'] >= min_experience)
    ]
    
    jd_requirements = extract_responsibilities(st.session_state.get('jd_text', ''))
    cv_text = st.session_state.get('cv_text', '')  # Retrieve cv_text from session state or set default
    matched = match_responsibilities(cv_text, jd_requirements)

    # Main results layout
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Total Candidates", len(df))
    with col2:
        st.metric("Average Score", f"{df['Score'].mean():.1f}")
    with col3:
        st.metric("Top Score", df['Score'].max())

    # Search bar
    search_term = st.text_input("Search candidates by name, skills, or keywords:")
    if search_term:
        filtered_df = filtered_df[
            filtered_df.apply(lambda row: any(
                search_term.lower() in str(row[col]).lower() 
                for col in ['Name', 'Skills', 'Certifications', 'Education', 'Responsibilities Matched']
            ), axis=1)
        ]

    # Interactive dataframe with all columns
    st.subheader("Ranked Candidates")
    def configure_dataframe():
        """Improved table formatting"""
        return {
            "Name": st.column_config.TextColumn(width="medium"),
            "Score": st.column_config.ProgressColumn(
                format="%d%%",
                min_value=0,
                max_value=100,
                width="small"
            ),
            "Experience": st.column_config.NumberColumn(
                format="%d yrs",
                width="small"
            ),
            "Skills": st.column_config.ListColumn(
                width="large",
                help="Top 10 relevant skills"
            ),
            "Certifications": st.column_config.ListColumn(
                width="medium",
                help="Professional certifications"
            ),
            "Education": st.column_config.ListColumn(
                width="large",
                help="Educational qualifications"
            ),
            "Responsibilities Matched": st.column_config.ListColumn(
                width="large",
                help="Top 5 matched responsibilities"
            )
        }

    # Candidate details expander
    st.subheader("Candidate Details")
    for idx, candidate in filtered_df.iterrows():
        with st.expander(f"{candidate['Name']} - Score: {candidate['Score']}%"):
            col1, col2 = st.columns([1, 3])
            
            with col1:
                st.markdown(f"**Email:** {candidate['Email']}")
                st.markdown(f"**Experience:** {candidate['Experience']} years")
                st.markdown(f"**Certifications:** {candidate['Certifications']}")
                st.markdown(f"**Education:** {candidate['Education']}")
                st.markdown(candidate['File'], unsafe_allow_html=True)
                
            with col2:
                st.markdown("**Matched Responsibilities**")
                st.write(candidate['Responsibilities Matched'])
                
                st.markdown("**Email Subject**")
                st.write(candidate['Subject'])

    # Download button
    st.download_button(
        label="Download Full Results CSV",
        data=filtered_df.drop(columns=["File"]).to_csv(index=False).encode("utf-8"),
        file_name="filtered_results.csv",
        mime="text/csv"
    )

    if st.button("Back to Main Page"):
        st.session_state.show_results = False
        st.rerun()
        
# Main app flow
if __name__ == "__main__":
    if 'show_results' not in st.session_state:
        st.session_state.show_results = False

    if st.session_state.show_results:
        results_page()
    else:
        main_page()


# # Refactored: Resume Analyzer with improved accuracy, modularity, and performance
# # Author: Intern (with help from ChatGPT)
# # Date: 2025-05

# import streamlit as st
# import re
# import os
# import io
# import base64
# import email
# import imaplib
# import datetime
# import hashlib
# import docx
# import PyPDF2
# import pandas as pd
# import spacy
# from fpdf import FPDF
# from email.header import decode_header
# from sklearn.feature_extraction.text import TfidfVectorizer
# from sklearn.metrics.pairwise import cosine_similarity
# from rapidfuzz import fuzz

# # Load NLP model
# try:
#     nlp = spacy.load("en_core_web_sm")
# except:
#     import subprocess
#     subprocess.run(["python", "-m", "spacy", "download", "en_core_web_sm"])
#     nlp = spacy.load("en_core_web_sm")

# # --- Utilities ---
# def decode_subject(subject):
#     decoded = decode_header(subject)
#     return ''.join([part.decode(enc or 'utf-8') if isinstance(part, bytes) else part for part, enc in decoded])

# def decode_filename(filename):
#     decoded = decode_header(filename)
#     return ''.join([part.decode(enc or 'utf-8') if isinstance(part, bytes) else part for part, enc in decoded])

# def extract_text_from_pdf(file_content):
#     try:
#         reader = PyPDF2.PdfReader(io.BytesIO(file_content))
#         return " ".join([page.extract_text() or "" for page in reader.pages])
#     except:
#         return ""

# def extract_text_from_docx(file_content):
#     try:
#         doc = docx.Document(io.BytesIO(file_content))
#         return " ".join(p.text for p in doc.paragraphs)
#     except:
#         return ""

# def clean_text(text):
#     text = text.replace('\n', ' ').replace('\r', ' ')
#     return re.sub(r'[^\x00-\x7F]+', ' ', text)

# def hash_text(text):
#     return hashlib.sha256(text.encode('utf-8')).hexdigest()

# # --- CV Extraction ---
# def extract_name(text):
#     doc = nlp(text[:1000])
#     for ent in doc.ents:
#         if ent.label_ == "PERSON":
#             return ent.text
#     return "Unknown"

# def extract_email(text):
#     match = re.search(r'[\w\.-]+@[\w\.-]+', text)
#     return match.group(0) if match else ""

# def extract_skills(text):
#     keywords = ["python", "sql", "aws", "azure", "java", "excel", "airflow", "kafka", "spark"]
#     text = text.lower()
#     skills = set()
#     for kw in keywords:
#         if kw in text:
#             skills.add(kw)
#     return list(skills)

# def extract_certifications(text):
#     patterns = [
#         r'aws certified [\w\s-]+',
#         r'certified [\w\s-]+',
#         r'pmp', r'ccna', r'cissp', r'azure fundamentals', r'cisa', r'microsoft certified.*?'
#     ]
#     found = set()
#     for pattern in patterns:
#         found.update(re.findall(pattern, text, re.I))
#     return list(found)

# def extract_education(text):
#     matches = re.findall(
#         r'(Bachelor|Master|PhD|MBA|BSc|MSc)[^\n]{0,100}?(in [^,\n]+)?[^\n]*?(at|from)?\s*([^\n,\r]+)?[,\s]*(\d{4})?[-–]?(\d{4}|Present)?',
#         text, re.IGNORECASE
#     )
#     results = []
#     for m in matches:
#         degree = m[0].strip()
#         course = m[1].strip() if m[1] else ''
#         school = m[3].strip() if m[3] else ''
#         start = m[4] if m[4] else ''
#         end = m[5] if m[5] else ''
#         entry = f"{degree} {course} from {school} ({start}–{end})"
#         results.append(entry.strip())
#     return results

# def extract_experience(text):
#     matches = re.findall(r'(\d{4}).*?(\d{4}|Present)', text)
#     years = []
#     for start, end in matches:
#         try:
#             start_year = int(start)
#             end_year = datetime.datetime.now().year if "present" in end.lower() else int(end)
#             years.append((start_year, end_year))
#         except:
#             continue
#     if not years:
#         return 0
#     start_years = [y[0] for y in years]
#     return datetime.datetime.now().year - min(start_years)

# def extract_responsibilities(jd_text):
#     matches = re.findall(r'(responsibilities|requirements|key tasks)[:\-\s]*([\s\S]+?)(?=(\n\n|\n[A-Z]))', jd_text, re.I)
#     if matches:
#         return [line.strip() for m in matches for line in m[1].split("\n") if 5 < len(line.strip()) < 150]
#     return []

# # --- PDF generation ---
# def export_cv_to_pdf(candidate):
#     pdf = FPDF()
#     pdf.add_page()
#     pdf.set_font("Arial", size=12)
#     pdf.multi_cell(0, 8, txt=f"Resume: {candidate['Name']}\n\n{candidate['RawText']}")
#     path = f"{candidate['Name'].replace(' ', '_')}_resume.pdf"
#     pdf.output(path)
#     return path

# # --- Email Functions ---
# def fetch_attachments(email_user, email_pass, keyword_filter, days_back, max_emails):
#     try:
#         server = "imap.gmail.com" if "gmail" in email_user else "imap.mail.yahoo.com" if "yahoo" in email_user else "imap-mail.outlook.com"
#         mail = imaplib.IMAP4_SSL(server)
#         mail.login(email_user, email_pass)
#         mail.select("inbox")
#         date_filter = (datetime.datetime.now() - datetime.timedelta(days=days_back)).strftime("%d-%b-%Y")
#         search = f'(SINCE {date_filter}' + (f' SUBJECT "{keyword_filter}"' if keyword_filter else '') + ')'
#         status, messages = mail.search(None, search)
#         ids = list(set(messages[0].split()))[-max_emails:]
#         attachments = []
#         for i in ids[::-1]:
#             _, msg_data = mail.fetch(i, "(RFC822)")
#             msg = email.message_from_bytes(msg_data[0][1])
#             subject = decode_subject(msg.get("Subject", ""))
#             for part in msg.walk():
#                 filename = part.get_filename()
#                 if filename:
#                     filename = decode_filename(filename)
#                     ext = os.path.splitext(filename)[1].lower()
#                     content = part.get_payload(decode=True)
#                     if ext in [".pdf", ".docx"]:
#                         attachments.append((filename, content, subject))
#         return attachments
#     except Exception as e:
#         st.error(f"Email error: {e}")
#         return []

# # --- Scoring ---
# def score_cv(cv_text, jd_text):
#     try:
#         vectorizer = TfidfVectorizer(stop_words='english')
#         tfidf = vectorizer.fit_transform([jd_text, cv_text])
#         similarity = cosine_similarity(tfidf[0:1], tfidf[1:2])[0][0]
#     except:
#         similarity = 0
#     return round(similarity * 100)

# # --- Streamlit UI ---
# st.title("📄 AI Resume Analyzer")
# email_user = st.text_input("Email Address")
# email_pass = st.text_input("Email Password (App Password)", type="password")
# subject_filter = st.text_input("Filter by Subject or Keyword (Optional)")
# days_back = st.slider("Only check emails from the past X days", min_value=1, max_value=90, value=30)
# max_emails = st.slider("Max emails to fetch", min_value=10, max_value=100, value=50)

# st.markdown("### Job Description")
# jd_text = st.text_area("Paste Job Description", height=200)
# uploaded_jd = st.file_uploader("Or upload Job Description (PDF/DOCX)", type=["pdf", "docx"])
# if uploaded_jd is not None:
#     jd_text = extract_text_from_pdf(uploaded_jd.read()) if uploaded_jd.name.endswith(".pdf") else extract_text_from_docx(uploaded_jd.read())

# jd_responsibilities = extract_responsibilities(jd_text)

# if st.button("Fetch & Analyze Resumes"):
#     if not email_user or not email_pass or not jd_text.strip():
#         st.error("Fill all fields.")
#     else:
#         with st.spinner("Processing resumes..."):
#             seen_hashes = set()
#             results = []
#             attachments = fetch_attachments(email_user, email_pass, subject_filter, days_back, max_emails)
#             for fname, content, subject in attachments:
#                 ext = os.path.splitext(fname)[1].lower()
#                 cv_text = extract_text_from_pdf(content) if ext == ".pdf" else extract_text_from_docx(content)
#                 cv_text = clean_text(cv_text)
#                 if len(cv_text) < 100:
#                     continue
#                 hash_val = hash_text(cv_text)
#                 if hash_val in seen_hashes:
#                     continue
#                 seen_hashes.add(hash_val)

#                 score = score_cv(cv_text, jd_text)
#                 name = extract_name(cv_text)
#                 email_addr = extract_email(cv_text)
#                 skills = extract_skills(cv_text)
#                 certs = extract_certifications(cv_text)
#                 edu = extract_education(cv_text)
#                 exp = extract_experience(cv_text)
#                 matched = [r for r in jd_responsibilities if r.lower() in cv_text.lower()]

#                 pdf_path = export_cv_to_pdf({"Name": name, "RawText": cv_text})
#                 with open(pdf_path, "rb") as f:
#                     encoded = base64.b64encode(f.read()).decode()
#                 download_link = f'<a href="data:application/pdf;base64,{encoded}" download="{pdf_path}">Download PDF</a>'

#                 results.append({
#                     "Name": name,
#                     "Email": email_addr,
#                     "Score": score,
#                     "Experience": exp,
#                     "Skills": ", ".join(skills),
#                     "Certifications": ", ".join(certs),
#                     "Education": ", ".join(edu),
#                     "Responsibilities Matched": ", ".join(matched),
#                     "File": download_link,
#                     "Subject": subject,
#                 })

#             if results:
#                 df = pd.DataFrame(results)
#                 st.write("### Ranked Candidates")
#                 st.write(df.to_html(escape=False), unsafe_allow_html=True)

#                 csv = df.drop(columns=["File"]).to_csv(index=False).encode("utf-8")
#                 st.download_button("Download Results CSV", csv, "ranked_results.csv")
#             else:
#                 st.warning("No valid resumes found.")